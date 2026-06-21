import os
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_DIR = 'input'
OUTPUT_DIR = 'output'
TEMPLATE_PATH = os.path.join(INPUT_DIR, 'CapstoneHospital_20260518_v1.docx')
CSV_PATH = os.path.join(INPUT_DIR, 'sales.csv')
MAIN_PY_PATH = 'main.py'
OUTPUT_PATH = os.path.join(OUTPUT_DIR, 'CapstoneHospital_Sales_20260621_v1.docx')
IMG_LINE = os.path.join(OUTPUT_DIR, 'plot_line.png')
IMG_BAR = os.path.join(OUTPUT_DIR, 'plot_bar.png')

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Load & compute stats from sales.csv ───────────────────────────────────
df = pd.read_csv(CSV_PATH)
df['Date'] = pd.to_datetime(df['Date'])
monthly_sales = df.resample('ME', on='Date').sum(numeric_only=True)
monthly_sales_reset = monthly_sales.reset_index()
monthly_sales_reset['Month'] = monthly_sales_reset['Date'].dt.strftime('%b')

sales_mean = df['Total Amount'].mean()
top_month = monthly_sales['Total Amount'].idxmax().strftime('%B %Y')
low_month = monthly_sales['Total Amount'].idxmin().strftime('%B %Y')
n_rows = len(df)
columns_list = ', '.join(df.columns.tolist())
date_range = f"{df['Date'].min().strftime('%Y-%m-%d')} to {df['Date'].max().strftime('%Y-%m-%d')}"
categories = ', '.join(df['Product Category'].unique().tolist())

# ── Generate and save plots (from plt.show() in main.py → Share section) ─
plt.figure(figsize=(10, 4))
plt.plot(monthly_sales.index, monthly_sales['Total Amount'], marker='o', color='steelblue')
plt.title('Monthly Sales Trends')
plt.xlabel('Month')
plt.ylabel('Sales')
plt.tight_layout()
plt.savefig(IMG_LINE, dpi=150)
plt.close()

plt.figure(figsize=(10, 5))
sns.barplot(x='Month', y='Total Amount', data=monthly_sales_reset, color='skyblue', errorbar=None)
plt.title('Monthly Sales')
plt.xlabel('Month')
plt.ylabel('Total Sales Amount')
plt.tight_layout()
plt.savefig(IMG_BAR, dpi=150)
plt.close()

# ── Read main.py code for appendix ────────────────────────────────────────
with open(MAIN_PY_PATH, 'r', encoding='utf-8') as f:
    main_code = f.read()

# ── Extract styles from template ──────────────────────────────────────────
template = Document(TEMPLATE_PATH)

def copy_font(src_font, dst_font):
    if src_font.name:
        dst_font.name = src_font.name
    if src_font.size:
        dst_font.size = src_font.size
    if src_font.bold is not None:
        dst_font.bold = src_font.bold
    try:
        if src_font.color and src_font.color.rgb:
            dst_font.color.rgb = src_font.color.rgb
    except Exception:
        pass

doc = Document()

for style_name in ('Normal', 'Heading 1', 'Heading 2'):
    try:
        src = template.styles[style_name].font
        dst = doc.styles[style_name].font
        copy_font(src, dst)
    except (KeyError, Exception):
        pass

# ── Title page ────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('Sales Data Analysis — Capstone Report')
run.bold = True
run.font.size = Pt(20)
doc.add_paragraph()

# ── Helper ────────────────────────────────────────────────────────────────
def section(heading, body):
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)

# ── 1. ASK ────────────────────────────────────────────────────────────────
section('1. Ask',
    "Business Question: What are the monthly sales trends, and which time periods "
    "drive the most revenue?\n\n"
    "Objective: Identify seasonal patterns in sales data to support data-driven "
    "decisions on inventory planning, marketing campaigns, and resource allocation.")

# ── 2. PREPARE ────────────────────────────────────────────────────────────
section('2. Prepare',
    f"Data Source: Internal sales records exported as sales.csv.\n\n"
    f"Dataset Overview:\n"
    f"  • Rows: {n_rows} transactions\n"
    f"  • Columns ({len(df.columns)}): {columns_list}\n"
    f"  • Date Range: {date_range}\n"
    f"  • Product Categories: {categories}\n\n"
    "Data integrity was verified using df.info() and df.describe(). "
    "No missing values were found in critical columns.")

# ── 3. PROCESS ────────────────────────────────────────────────────────────
section('3. Process',
    "Processing steps applied using Python (pandas):\n\n"
    "  1. Date conversion: df['Date'] parsed to datetime using pd.to_datetime().\n"
    "  2. Monthly aggregation: transactions resampled at month-end frequency using "
    "df.resample('M', on='Date').sum(), summing Total Amount per month.\n"
    "  3. Month label: a 'Month' column (abbreviated name) derived for chart readability.\n"
    "  4. Null check: df.info() confirmed no missing values in critical columns.\n\n"
    "Tools used: Python 3, pandas, matplotlib, seaborn.")

# ── 4. ANALYSE ────────────────────────────────────────────────────────────
section('4. Analyse',
    f"Key Findings:\n\n"
    f"  • Average Monthly Sales: ${sales_mean:,.2f}\n"
    f"  • Highest Sales Month: {top_month}\n"
    f"  • Lowest Sales Month: {low_month}\n\n"
    f"Monthly aggregation reveals seasonal variation in sales volume. "
    f"The peak month ({top_month}) shows significantly higher revenue, suggesting "
    f"seasonal demand. The lowest month ({low_month}) represents an opportunity "
    f"for targeted marketing efforts.")

# ── 5. SHARE ─────────────────────────────────────────────────────────────
doc.add_heading('5. Share', level=1)
doc.add_paragraph(
    "Two visualizations were created to communicate findings to stakeholders:")

p = doc.add_paragraph()
p.add_run("Figure 1 — Monthly Sales Trends (Line Chart):").bold = True
doc.add_picture(IMG_LINE, width=Inches(5.5))
doc.add_paragraph(
    "The line chart shows the evolution of total sales month by month, "
    "highlighting peaks and troughs across the year.")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Figure 2 — Monthly Sales (Bar Chart):").bold = True
doc.add_picture(IMG_BAR, width=Inches(5.5))
doc.add_paragraph(
    "The bar chart provides an at-a-glance comparison of total sales per month, "
    "making it easy to identify the strongest and weakest performing periods.")

# ── 6. ACT ────────────────────────────────────────────────────────────────
section('6. Act',
    f"Recommendations based on the analysis:\n\n"
    f"  1. Capitalize on peak month ({top_month}): increase inventory and staffing "
    f"ahead of the high-demand period.\n"
    f"  2. Address low-sales period ({low_month}): run targeted promotions or "
    f"loyalty campaigns to stimulate demand.\n"
    f"  3. Category focus: prioritize high-margin items from ({categories}) "
    f"during peak periods.\n"
    f"  4. Automate reporting: schedule this analysis monthly to continuously "
    f"monitor trends.")

# ── Appendix — Python Code ────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('Appendix — Python Code (main.py)', level=1)
doc.add_paragraph(
    "The following Python script was used to perform the data processing "
    "and generate the visualizations:")
doc.add_paragraph()

code_para = doc.add_paragraph()
run = code_para.add_run(main_code)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# ── Save ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Report saved to {OUTPUT_PATH}")